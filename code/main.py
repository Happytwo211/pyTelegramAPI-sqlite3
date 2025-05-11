import telebot
import datetime
from docx import Document
import sqlite3
from keyboards import (start_keyboard, departure_cities_keyboard, arrival_cities_keyboard, adult_tourists, kids_tourists,
                            stars_hotel, dinner_quantity, days_long, peiod, yes_or_no_keyboard,
                            user_contancts, get_phone, consult_kb, test)

from token_ import TOKEN

connection = sqlite3.connect('../data/db_ALBA_1.sqlite3', check_same_thread=False)
cursor = connection.cursor()


bot = telebot.TeleBot(TOKEN)

user_data = []




@bot.message_handler(commands=['help', 'start'])
def handle_start(message):

    bot.send_message(message.chat.id, f'Привет, {message.from_user.first_name}'
                                      f'\nЭтот бот поможет тебе подобрать  тур!', reply_markup=start_keyboard())
    global user_name
    user_name = f'@{message.from_user.username}'

    return user_name

@bot.message_handler(commands=['db'])
def handle_db(message):
    if message.from_user.id == //id or message.from_user.id == //id:
        try:
            cursor.execute('''
            SELECT * FROM users_data
            ''')
            data = cursor.fetchall()
            if data:
                bot.send_message(message.chat.id, f'<b>Данные из основной базы данных!</b>\n\n', parse_mode='HTML')
                for i in data:
                    if i[11] is not None:
                        user_phone = f'{i[11]}'
                    else:
                        user_phone = f'Не указан'

                    bot.send_message(message.chat.id, f'Телеграм пользователя - {i[0]},'
                                                      f'\nГород вылета - {i[1]},'
                                                      f'\nГород прилета - {i[2]}'
                                                      f'\nКоличество взрослых - {i[3]},'
                                                      f'\nКоличесвто детей - {i[4]},'
                                                      f'\nКоличесвто звезд - {i[5]},'
                                                      f'\nТип питания - {i[6]},'
                                                      f'\nКоличесвто ночей - {i[7]},'
                                                      f'\nПериод отдыха - {i[8]}'
                                                      f'\nТелефон- {user_phone}')

            else:
                bot.send_message(message.chat.id, f'База данных пуста!')

        except sqlite3.OperationalError:
            bot.send_message(message.chat.id, f'Произошла  ошибка!')

    else:
        bot.send_message(message.chat.id, f'У вас нет доступа к этой команде')

@bot.message_handler(commands=['db_consult'])
def handle_db(message):
    if message.from_user.id == //id or message.from_user.id == /id:
        try:
            cursor.execute('''
            SELECT * FROM users_consult
            ''')
            data = cursor.fetchall()
            if data:
                bot.send_message(message.chat.id, f'Данные из базы данных по консультации\n\n')
                for i in data:
                    bot.send_message(message.chat.id,
                                                 f'\nTelegram пользователя - {i[2]},'
                                                      f'\nТелефон пользователя - {i[3]}')
            else:
                bot.send_message(message.chat.id, f'База данных пуста!')
        except sqlite3.OperationalError:
            bot.send_message(message.chat.id, f'Произошла ошибка! ')

    else:
        bot.send_message(message.chat.id, f'У вас нет доступа к этой команде')
@bot.message_handler(commands=['db_clear'])
def handle_db(message):
    if message.from_user.id == //id or message.from_user.id == //id :
        bot.send_message(message.chat.id, f'Вы хотите удалить основную базу данных?')
        bot.register_next_step_handler(message, delete_DB)

    else:
        bot.send_message(message.chat.id, f'У вас нет доступа к этой команде')

def delete_DB(message):
    if message.text.lower() == 'да':
        try:
            cursor.execute('''
            DELETE FROM users_data
            ''')
            connection.commit()
            bot.send_message(message.chat.id, f'основная БД очищена')
        except sqlite3.OperationalError:
            bot.send_message(message.chat.id, f'Произошла ошибка! ')
    else:
        bot.send_message(message.chat.id, f'Данные остались без изменения')

@bot.message_handler(commands=['db_consult_clear'])
def handle_db(message):
    if message.from_user.id == //id  or message.from_user.id == //id :
        bot.send_message(message.chat.id, f'Вы уверены, что хотите удалить БД консультаций?')
        bot.register_next_step_handler(message, delete_consult_DB)
    else:
        bot.send_message(message.chat.id, f'У вас нет доступа к этой команде')
@bot.message_handler(commands=['db_backup'])
def hanlde_backup(message):
    if message.from_user.id == //id  or message.from_user.id == //id :
        name = 'backup'
        name_count = 0
        document = Document()
        try:
            # Извлекаем данные из базы данных
            cursor.execute('SELECT * FROM users_data')
            rows = cursor.fetchall()
            # Создаем новый документ Word
            # Добавляем заголовки таблицы
            table = document.add_table(rows=1, cols=len(rows[0]))
            hdr_cells = table.rows[0].cells
            for i, col_header in enumerate(['user_name', 'departure_city', 'arrive_city',
                                            'adult_persons', 'children', 'stars',
                                            'eat_plan', 'days', 'period', 'user_tg', 'user_phone']):
                hdr_cells[i].text = col_header

            # Заполняем таблицу данными
            for row in rows:
                cells = table.add_row().cells
                for j, cell_value in enumerate(row):
                    cells[j].text = str(cell_value)


        except Exception as e:

            bot.send_message(message.chat.id, f'Произошла ошибка: {e}')
        try:
            # Сохраняем документ
            # document.save(fr'C:\Users\thsim\Documents\{name}.docx')
            document.save(f'../data/{name}.docx')
            # Сообщаем пользователю о сохранении данных
            bot.send_message(message.chat.id, f'Данные сохранены в файл {name}.docx.')

        except PermissionError:

            name_count = name_count + 1
            '_'.join(f'{name}')
            print(name)
            document.save(fr'C:\Users\thsim\Documents\{name}.docx')
            return name_count

        except Exception as e:
            bot.send_message(message.chat.id, f'Произошла ошибка: {e}')

    else:
        bot.send_message(message.chat.id, f'У вас нет доступа к этой команде')

@bot.message_handler(commands=['db_consult_backup'])
def hanlde_backup(message):
    if message.from_user.id == //id  or message.from_user.id == //id :
        name = 'backup_consult'
        name_count = 0
        document = Document()
        try:
            # Извлекаем данные из базы данных
            cursor.execute('SELECT * FROM users_consult')
            rows = cursor.fetchall()
            # Создаем новый документ Word
            document = Document()
            # Добавляем заголовки таблицы
            table = document.add_table(rows=1, cols=len(rows[0]))
            hdr_cells = table.rows[0].cells
            for i, col_header in enumerate(['user_name', 'telegram', 'phone_number']):
                hdr_cells[i].text = col_header

            # Заполняем таблицу данными
            for row in rows:
                cells = table.add_row().cells
                for j, cell_value in enumerate(row):
                    cells[j].text = str(cell_value)

        except Exception as e:
            bot.send_message(message.chat.id, f'Произошла ошибка: {e}')
        try:
            # Сохраняем документ
            # document.save(fr'C:\Users\thsim\Documents\{name}.docx')
            document.save(f'../data/{name}.docx')
            # Сообщаем пользователю о сохранении данных
            bot.send_message(message.chat.id, f'Данные сохранены в файл {name}.docx.')

        except PermissionError:

            name_count = name_count + 1
            '_'.join(f'{name}')
            print(name)
            document.save(fr'C:\Users\thsim\Documents\{name}.docx')
            return name_count

        except Exception as e:
            bot.send_message(message.chat.id, f'Произошла ошибка: {e}')
    else:
        bot.send_message(message.chat.id, f'У вас нет доступа к этой команде')
#calldata
def delete_consult_DB(message):
    if message.text.lower() == 'да':
        try:
            cursor.execute('''
            DELETE FROM users_consult
            ''')
            connection.commit()
            bot.send_message(message.chat.id, f'консалт БД очищена')
        except sqlite3.OperationalError:
            bot.send_message(message.chat.id, f'Произошла ошибка! ')
    else:
        bot.send_message(message.chat.id, f'Данные остались без изменения')

@bot.callback_query_handler(func=lambda call: call.data in ['phone_consult'])
def handle_consult_phone(call):
    message = call.message
    bot.send_message(message.chat.id, f'Поделитесь номером телефона и мы свяжемся с вами!', reply_markup= get_phone())
    bot.register_next_step_handler(message, phone_consult)

def phone_consult(message):
    try:
        phone_number_ = message.contact.phone_number
    except AttributeError:
        phone_number_ = message.text
    bot.send_message(message.chat.id, f'Скоро с вами свяжется специалист по номеру '
                                      f'\n{phone_number_}')
    try:
        cursor.execute('INSERT INTO users_consult (user_name, phone_number) VALUES (?, ?)', (user_name, phone_number_))
        connection.commit()
    except NameError:
        bot.send_message(message.chat.id, f'Произошла ошибка')


@bot.callback_query_handler(func=lambda call: call.data in ['telegram_consult'])
def handle_tg_consult(call):
    message = call.message
    username_tg_ = message.from_user.username

    try:
        cursor.execute('INSERT INTO users_consult (user_name, telegram) VALUES (?, ?)', (user_name, username_tg_))
        connection.commit()
        bot.edit_message_text(f'<strong>Специалист свяжется с вами в ближайшее время в telegram</strong>',
                              call.message.chat.id, call.message.message_id, parse_mode='HTML')
    except NameError:
        bot.send_message(message.chat.id, f'Произишла ошибка')


@bot.callback_query_handler(func=lambda call: call.data in ['get_tour'])
def handle_callback_get_tour(call):
    message = call.message
    bot.send_message(message.chat.id, f'Ваш город вылета?'
                                      f'\n█▒▒▒▒▒▒▒▒▒ 12%', reply_markup=departure_cities_keyboard())

@bot.callback_query_handler(func=lambda call: call.data in ['get_consult'])
def handle_callback_get_tour(call):
    message = call.message
    bot.send_message(message.chat.id, f'Как бы вы хотели получить консультацию?', reply_markup=consult_kb())


#CALLBACK_DEPARTURE_CITIES
@bot.callback_query_handler(func=lambda call: call.data in ['Москва', 'Санкт-Петербург', 'Екатеренбург',
                                                            'Новосибирск', 'Казань', 'Другое'])
def handle_callback_cities(call):
    bot.edit_message_text(
        f'В какой стране вы хотите отдохнуть? \n ██▒▒▒▒▒▒▒▒ 25%',
        call.message.chat.id, call.message.message_id, reply_markup=arrival_cities_keyboard()
    )
    global departure_city
    departure_city = call.data
    user_data.append(f'Город вылета - {departure_city}')
    return departure_city


@bot.callback_query_handler(func=lambda call: call.data in ['main_menu_arrival_cities'])
def handle_callback_departure_main_menu_(call):
    bot.edit_message_text(
        f'Ваш город вылета?'
        f'\n█▒▒▒▒▒▒▒▒▒ 12%', call.message.chat.id, call.message.message_id,
        reply_markup=departure_cities_keyboard()
    )
    try:
        user_data.pop()
    except IndexError:
        pass

#CALLDATA_ARRIVAL_CITIES
@bot.callback_query_handler(func=lambda call: call.data in ['Турция', 'Египет', 'ОАЭ',
                                                            'Россия', 'Таиланд', 'Китай',
                                                            'Мальдивы', 'Сейшелы'])
def handle_callback_arrival_cities(call):
    bot.edit_message_text(
        'Выберите количество взрослых туристов\n███▒▒▒▒▒▒▒ 38%',
        call.message.chat.id, call.message.message_id, reply_markup=adult_tourists()

    )
    global arrive_city
    arrive_city = call.data
    user_data.append(f'Город прилета  -{arrive_city}')
    return arrive_city

@bot.callback_query_handler(func=lambda call: call.data in ['main_menu_adult_tourist'])
def handle_main_menu_adult_tourist(call):
    bot.edit_message_text(
        f'В какой стране вы хотите отдохнуть? \n ██▒▒▒▒▒▒▒▒ 25%',
        call.message.chat.id, call.message.message_id, reply_markup=arrival_cities_keyboard()
    )
    try:
        user_data.pop()
    except IndexError:
        pass


#CALLBACK_ADULTS_TOURIST
@bot.callback_query_handler(func=lambda call: call.data in ['👤', '👤👤', '👤👤👤', '👤👤👤👤', '👤👤👤👤👤', 'Другое количество' ])
def handle_adult_tourist(call):
    bot.edit_message_text(
        'Выберите количество детей\n█████▒▒▒▒▒ 50%',
        call.message.chat.id, call.message.message_id, reply_markup=kids_tourists()

    )
    global adult_persons
    adult_persons = call.data
    user_data.append(f'Количество взрослых - {adult_persons}')
    return adult_persons

@bot.callback_query_handler(func=lambda call: call.data in ['main_menu_kids_tourist'])
def handle_main_menu_adult_tourist(call):
    bot.edit_message_text(
        'Выберите количество взрослых туристов\n███▒▒▒▒▒▒▒ 38%',
        call.message.chat.id, call.message.message_id, reply_markup=adult_tourists()

    )
    try:
        user_data.pop()
    except IndexError:
        pass



#CALLBACK_KIDS_TOURIST
@bot.callback_query_handler(func=lambda call: call.data in ['👶', '👶👶', '👶👶👶', '👶👶👶👶', 'без детей', 'Другое количество дети' ])
def handle_adult_tourist(call):
    bot.edit_message_text(
        'Выберите количество звезд отеля\n██████▒▒▒▒ 62%',
        call.message.chat.id, call.message.message_id, reply_markup=stars_hotel()

    )
    global children
    children = call.data
    user_data.append(f'Количество детей - {children}')
    return children

@bot.callback_query_handler(func=lambda call: call.data in ['main_menu_stars_hotel'])
def handle_adult_tourist(call):
    bot.edit_message_text(
        'Выберите количество детей\n█████▒▒▒▒▒ 50%',
        call.message.chat.id, call.message.message_id, reply_markup=kids_tourists()

    )
    try:
        user_data.pop()
    except IndexError:
        pass

#CALLBACK_DINERS
@bot.callback_query_handler(func=lambda call: call.data in ['⭐️', '⭐️⭐️', '⭐️⭐️⭐️', '⭐️⭐️⭐️⭐️', '⭐️⭐️⭐️⭐️⭐️'])
def handle_adult_tourist(call):
    bot.edit_message_text(
        'Выберите тип питания \n███████▒▒▒ 75%',
        call.message.chat.id, call.message.message_id, reply_markup=dinner_quantity()

    )
    global stars
    stars = call.data
    user_data.append(f'Количество звезд - {stars}')
    return stars

@bot.callback_query_handler(func=lambda call: call.data in ['main_menu_dinner'])
def handle_adult_tourist(call):
    bot.edit_message_text(
        'Выберите количество звезд отеля\n██████▒▒▒▒ 62%',
        call.message.chat.id, call.message.message_id, reply_markup=stars_hotel()

    )
    try:
        user_data.pop()
    except IndexError:
        pass


#CAllBACK DINERS

@bot.callback_query_handler(func=lambda call: call.data in ['Все включено', 'Завтрак', 'Завтрак и ужин', 'Завтрак, обед и ужин'])
def handle_days_long(call):
    bot.edit_message_text(
        'Выберите количество ночей\n█████████▒  88%',
        call.message.chat.id, call.message.message_id, reply_markup=days_long()

    )
    global eat_plan
    eat_plan = call.data
    user_data.append(f'Тип питания -{eat_plan}')
    return eat_plan

@bot.callback_query_handler(func=lambda call: call.data in ['main_menu_days_long'])
def handle_adult_tourist(call):
    bot.edit_message_text(
        'Выберите Тип питания \n███████▒▒▒ 75%',
        call.message.chat.id, call.message.message_id, reply_markup=dinner_quantity()

    )
    try:
        user_data.pop()
    except IndexError:
        pass


#CALLBACK DAYS LONG
@bot.callback_query_handler(func=lambda call: call.data in ['7-9', '10-12','13-15'])
def handle_days_long(call):
    bot.edit_message_text(
        'В какой период вы планируете поездку?\n██████████ 100%',
        call.message.chat.id, call.message.message_id, reply_markup=peiod()

    )
    global days
    days = call.data
    user_data.append(f'Количество ночей -{days}')
    return days

@bot.callback_query_handler(func=lambda call: call.data in ['main_menu_days_long'])
def handle_days_long(call):
    bot.edit_message_text(
        'Выберите количесвто ночей\n█████████▒  88%',
        call.message.chat.id, call.message.message_id, reply_markup=days_long()

    )
    try:
        user_data.pop()
    except IndexError:
        pass

#CALLBACK PEDIOD

@bot.callback_query_handler(func=lambda call: call.data in ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь'])
def handle_days_long(call):
    bot.edit_message_text(
        f'Все готово!'
        f'\nПроверьте правильность заполненных данных',
        call.message.chat.id, call.message.message_id

    )

    global peiod_
    peiod_= call.data
    user_data.append(f'Период отдыха - {peiod_}')


    message = call.message
    for data in user_data:
        bot.send_message(message.chat.id, f'{data}')

    user_data.clear()

    bot.send_message(message.chat.id, f'Данные верны?', reply_markup=yes_or_no_keyboard())
    return peiod_
@bot.callback_query_handler(func=lambda call: call.data in ['main_menu_period'])
def handle_days_long(call):
    bot.edit_message_text(
        'В какой период вы планируете поездку?\n██████████ 100%',
        call.message.chat.id, call.message.message_id, reply_markup=peiod()

    )
    try:
        user_data.pop()
    except IndexError:
        pass


@bot.callback_query_handler(func=lambda call: call.data in ['Данные верны'])
def correct(call):
    message = call.message
    global message_id
    message_id = bot.edit_message_text(
        'Как можно с вами связаться?',
        call.message.chat.id, call.message.message_id, reply_markup=user_contancts()
    )


@bot.callback_query_handler(func=lambda call: call.data in ['Заполнить заного'])
def incorrect(call):
    message = call.message
    bot.edit_message_text(f'Привет, {message.from_user.first_name}'
                                      f'\nЭтот бот поможет тебе подобрать тебе тур!', call.message.chat.id, call.message.message_id,
                          reply_markup=start_keyboard())



user_tg_contacts=[]

@bot.callback_query_handler(func=lambda call: call.data in ['telegram'])
def telegram(call):
    message = call.message
    username_tg_ = message.from_user.username
    user_tg_contacts.append(f'Ник в телеграмме - {username_tg_}')
    try:
        cursor.execute('INSERT INTO users_data (user_name, departure_city, arrive_city, '
                       'adult_persons, children, stars, eat_plan, days, period, user_tg) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                       (user_name, departure_city, arrive_city,
                        adult_persons, children, stars,
                        eat_plan, days, peiod_, username_tg_))
        connection.commit()
        bot.edit_message_text(f'<strong>В ближайшее время с вами свяжется специалист в telegram</strong>',
                              call.message.chat.id, call.message.message_id, parse_mode='HTML')
        channel_id = -1002503000438
        bot.send_message(channel_id, f'В дату базу внесена новая записать\n'
                                     f'{datetime.date.today()}')
    except NameError:
        bot.send_message(message.chat.id, f'Произишла ошибка')



    user_tg_contacts.clear()




@bot.callback_query_handler(func=lambda call: call.data in ['number'])
def phone_num(call):
    message = call.message
    msg = bot.send_message(message.chat.id, f'\nПоделитесь номером телефона!', reply_markup=get_phone())

    bot.register_next_step_handler(msg, phone_number)
def phone_number(message):
    try:
        phone_number_ = message.contact.phone_number
    except AttributeError:
        phone_number_ = message.text
    bot.send_message(message.chat.id, f'Скоро с вами свяжется специалист по номеру '
                                      f'\n{phone_number_}')
    try:
        cursor.execute('INSERT INTO users_data (user_name, departure_city, arrive_city, '
                       'adult_persons, children, stars, eat_plan, days, period, user_phone) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                       (user_name, departure_city, arrive_city,
                        adult_persons, children, stars,
                        eat_plan, days, peiod_, phone_number_))
        connection.commit()
        channel_id = -1002503000438
        bot.send_message(channel_id, f'В дату базу внесена новая записать\n'
                                     f'{datetime.date.today()}')
    except NameError:
        bot.send_message(message.chat.id, f'Произошла ошибка!')







# connection.commit()
# connection.close()


if __name__ == "__main__":
    bot.polling(none_stop=True)

